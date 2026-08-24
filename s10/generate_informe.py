"""
Actualiza el informe S10 a partir del Word S8 (mismo formato),
corrigiendo métricas, anexo y secciones pendientes según el repositorio.

Ejecutar desde la raíz:
  python s10/organize_evidencias.py
  python s10/capture_demo.py
  python s10/generate_informe.py
"""
from __future__ import annotations

import json
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
S8_TEMPLATE = ROOT / "s8" / "docs" / "S8_Informe_Final.docx"
DOCS = Path(__file__).resolve().parent / "docs"
EVID = Path(__file__).resolve().parent / "evidencias"
CAP = EVID / "capturas"
OUT_DOCX = DOCS / "S10_Avance_Consolidado.docx"
OUT_PDF = DOCS / "S10_Avance_Consolidado.pdf"
REPO_URL = "https://github.com/marcoz1691/MIA_Deteccion_HC"


def _load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def _fmt(x: float, nd: int = 3) -> str:
    return f"{x:.{nd}f}".replace(".", ",")


def _pct(x: float) -> str:
    return f"{x * 100:.1f} %".replace(".", ",")


def load_metrics() -> dict:
    tfidf = _load_json(ROOT / "s6" / "metricas_ajuste.json") or _load_json(EVID / "metricas_tfidf.json")
    return {
        "tfidf": tfidf,
        "tripartita": _load_json(ROOT / "salidas_s7" / "metricas_tripartita.json")
        or _load_json(EVID / "metricas_tripartita.json"),
        "por_tipo": _load_json(ROOT / "salidas_s7" / "analisis_por_tipo.json")
        or _load_json(EVID / "analisis_por_tipo.json"),
        "idioma": _load_json(ROOT / "salidas_s7" / "eval_idioma_en_es.json")
        or _load_json(EVID / "eval_idioma_en_es.json"),
        "tfidf_idioma": _load_json(ROOT / "salidas_s7" / "eval_tfidf_idioma.json")
        or _load_json(EVID / "eval_tfidf_idioma.json"),
        "citimed_piloto": _load_json(ROOT / "salidas_s7" / "prueba_citimed.json"),
        "eval_citimed": _load_json(ROOT / "salidas_s7" / "eval_citimed.json"),
    }


def _citimed_piloto_text(citimed: dict, eval_c: dict) -> dict:
    """Resume métricas del piloto CITIMED (avance, no corpus final)."""
    ev = eval_c.get("metricas") or {}
    inferencias = citimed.get("inferencias_piloto") or citimed.get("inferencias") or []
    n_hist = len(citimed.get("anonimizados") or [])
    n_oraciones_eval = eval_c.get("n_oraciones")
    tiene_eval = bool(ev.get("roc_auc") is not None and n_oraciones_eval)
    return {
        "n_historias": n_hist or (1 if inferencias else 0),
        "n_oraciones_eval": n_oraciones_eval,
        "n_oraciones_inferencia": sum(i.get("n_oraciones", 0) for i in inferencias),
        "roc": ev.get("roc_auc") if tiene_eval else None,
        "auprc": ev.get("auprc") if tiene_eval else None,
        "modo": eval_c.get("modo", "cross_domain"),
        "entrenamiento": eval_c.get("entrenamiento", "MEDEC_train"),
        "tiene_piloto": bool(inferencias or citimed.get("anonimizados")),
        "tiene_eval": tiene_eval,
    }


def _fmt_opt(x: float | None, nd: int = 3) -> str:
    return _fmt(x, nd) if x is not None else "—"


def _bloque_verificacion_s10() -> list[tuple[str, str]]:
    """§3.4 — FastAPI, utilidades CITIMED, tests y .gitignore."""
    return [
        ("Heading 2", "3.4. Verificación FastAPI y calidad del código (S10)"),
        (
            "Normal",
            "Verificación operativa del backend (23-ago-2026): con "
            "`uvicorn api.main:app --port 8000`, GET /health respondió status=ok y "
            "modelo_tfidf_disponible=true (salidas_ajuste/modelo_ajustado.joblib). "
            "POST /generar con nota de alergia a penicilina y prescripción de amoxicilina "
            "(mock_llm=true) devolvió HTTP 200; top1=\"Se indica amoxicilina…\", alerta=true, "
            "brazos tfidf + llm_zero + llm_rag activos. Documentación interactiva en /docs. "
            "Frontend React (frontend/, puerto 5173) consume la misma API vía proxy Vite.",
        ),
        (
            "Normal",
            "Refactor del pipeline CITIMED (code review S10): s10/citimed_utils.py centraliza "
            "extracción del cuerpo clínico, detección heurística de PHI residual y lectura de "
            "salidas anonimizadas (.txt/.pdf). s10/probar_citimed.py separa eval cross-domain "
            "(solo plantilla etiquetada) de inferencias_piloto; aborta si queda PHI (--force "
            "solo depuración). demo/ejemplos.py reutiliza citimed_utils. Tests automatizados: "
            "s10/test_citimed_pipeline.py (6 casos pytest: PHI, CSV, cabeceras).",
        ),
        (
            "Normal",
            "Controles de repositorio (.gitignore): data/citimed_odontologia.csv generado, "
            "s10/anonimizador/salidas/, historia_ejemplo.txt con PHI aparente (uso local); "
            "en git solo ejemplos sintéticos (ejemplos/historia_ejemplo.sintetico.txt) y salidas "
            "anonimizadas de demostración. Entrega S10: consolidación en rama main pendiente de "
            "commit/push final con anonimizador, cleanup CITIMED e informe regenerado.",
        ),
    ]


def _insert_paragraphs_after(doc, anchor_contains: str, blocks: list[tuple[str, str]]):
    """Inserta párrafos después de un ancla. blocks = [(estilo, texto), ...]. Retorna el último párrafo."""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    anchor = _find_paragraph(doc, anchor_contains)
    if anchor is None:
        print(f"[warn] No se insertó bloque; ancla no encontrada: {anchor_contains[:50]}")
        return None
    prev = anchor._element
    last_para = anchor
    for style_name, text in blocks:
        new_p = OxmlElement("w:p")
        prev.addnext(new_p)
        para = Paragraph(new_p, anchor._parent)
        if style_name:
            try:
                para.style = style_name
            except Exception:
                pass
        para.add_run(text)
        prev = new_p
        last_para = para
    return last_para


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _find_paragraph(doc, contains: str):
    for p in doc.paragraphs:
        if contains in p.text:
            return p
    return None


def _find_table_by_header(doc, header: str):
    for tbl in doc.tables:
        if tbl.rows and header in tbl.rows[0].cells[0].text:
            return tbl
    return None


def _insert_picture_after(paragraph, image_path: Path, width_inches: float = 5.8, caption: str = "") -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph

    if not image_path.exists():
        return
    new_p = paragraph.insert_paragraph_before("")  # placeholder; we'll use addnext via oxml
    # insert_paragraph_after pattern
    new_el = paragraph._element.addnext(paragraph._element.__class__())
    pic_para = Paragraph(new_el, paragraph._parent)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap_el = new_el.addnext(paragraph._element.__class__())
        cap_para = Paragraph(cap_el, paragraph._parent)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True


def _insert_picture_after_paragraph(paragraph, image_path: Path, caption: str, width: float = 5.8) -> None:
    """Inserta figura inmediatamente después del párrafo dado."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph

    if not image_path.exists():
        print(f"[warn] Figura no encontrada: {image_path}")
        return

    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    pic_para = Paragraph(new_p, paragraph._parent)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Inches(width))

    cap_p = OxmlElement("w:p")
    new_p.addnext(cap_p)
    cap_para = Paragraph(cap_p, paragraph._parent)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run(caption)
    run.italic = True


def _insert_picture_after_simple(doc, anchor_text: str, image_path: Path, caption: str, width: float = 5.8) -> None:
    """Inserta figura después del párrafo que contiene anchor_text."""
    anchor = _find_paragraph(doc, anchor_text)
    if anchor is None:
        print(f"[warn] Ancla no encontrada: {anchor_text[:40]}")
        return
    _insert_picture_after_paragraph(anchor, image_path, caption, width)


def _add_row(table, cells: list[str]) -> None:
    row = table.add_row()
    for i, val in enumerate(cells):
        if i < len(row.cells):
            row.cells[i].text = val


def update_from_s8(m: dict) -> None:
    from docx import Document

    if not S8_TEMPLATE.exists():
        raise FileNotFoundError(f"Plantilla S8 no encontrada: {S8_TEMPLATE}")

    DOCS.mkdir(parents=True, exist_ok=True)
    shutil.copy2(S8_TEMPLATE, OUT_DOCX)
    doc = Document(str(OUT_DOCX))

    prueba = m["tfidf"].get("prueba", {})
    roc = prueba.get("roc_auc", 0.9485)
    auprc = prueba.get("auprc", 0.4186)
    roc_ci = prueba.get("roc_auc_ci95", [0, 0.9398, 0.9576])
    auprc_ci = prueba.get("auprc_ci95", [0, 0.392, 0.446])
    loc = m["tfidf"].get("localizacion_top1_test", 0.8457)
    n_notas = m["tfidf"].get("notas_con_error_test", 311)
    aciertos = int(round(loc * n_notas))
    cv_mean = m["tfidf"].get("cv_auc_mean", 0.965)
    cv_sd = m["tfidf"].get("cv_auc_sd", 0.0072)

    trip = m["tripartita"].get("brazos", {})
    tf_t = trip.get("tfidf", {}).get("oracion", {})
    llm_t = trip.get("llm_zero", {}).get("oracion", {})
    rag_t = trip.get("llm_rag", {}).get("oracion", {})
    mock = m["tripartita"].get("llm_stats", {}).get("mock_mode", True)
    n_trip = m["tripartita"].get("n_oraciones", 500)
    mcnemar_p = m["tripartita"].get("mcnemar", {}).get("tfidf_vs_llm_zero", {}).get("p_value", 0.11)

    idioma = m["idioma"]
    delta_auc = idioma.get("delta_auc_en_minus_es", 0)
    tfidf_es = m.get("tfidf_idioma", {}).get("spanish", {}).get("test", {})
    piloto = _citimed_piloto_text(m.get("citimed_piloto", {}), m.get("eval_citimed", {}))
    if piloto["tiene_eval"]:
        citimed_eval_txt = (
            f"eval cross-domain sobre plantilla odontológica etiquetada "
            f"(n={piloto['n_oraciones_eval']}): ROC-AUC {_fmt_opt(piloto['roc'])}, "
            f"AUPRC {_fmt_opt(piloto['auprc'])}"
        )
    else:
        citimed_eval_txt = "pipeline piloto CITIMED (eval cross-domain pendiente de ejecutar)"
    piloto_infer_txt = (
        f"inferencia piloto sobre {piloto['n_historias']} historia(s) anonimizada(s) "
        f"({piloto['n_oraciones_inferencia']} oraciones, sin mezclar en métricas AUC)"
    )

    # --- Portada / metadatos ---
    for p in doc.paragraphs:
        if "Tarea S8" in p.text:
            _replace_paragraph_text(p, "S10 — Avance Consolidado del Proyecto y Documento Final")
        elif p.text.startswith("Fecha:"):
            _replace_paragraph_text(p, f"Fecha: {date.today().strftime('%d de %B de %Y').replace('August', 'agosto')}")
        elif p.text.startswith("Grupo:"):
            _replace_paragraph_text(
                p,
                "Grupo: Patricio Bayas Meza, José Puebla Paladines, Marco Zurita Rojas",
            )

    # --- Resumen ---
    resumen = _find_paragraph(doc, "Las historias clínicas contienen inconsistencias")
    if resumen:
        _replace_paragraph_text(
            resumen,
            "Las historias clínicas contienen inconsistencias —medicación contraindicada, incoherencias "
            "entre diagnóstico y tratamiento, contradicciones internas— que comprometen la seguridad del "
            "paciente. Este informe consolida el avance del proyecto MIA_Deteccion_HC para CITIMED: "
            "asistente con TF-IDF, LLM zero-shot y LLM+RAG evaluados sobre MEDEC. El baseline léxico a "
            "nivel de nota no supera el azar (ROC-AUC 0,504). Al reformular la detección a nivel de "
            f"oración, el modelo ajustado alcanza ROC-AUC {_fmt(roc)} (IC 95 % {_fmt(roc_ci[1])}–"
            f"{_fmt(roc_ci[2])}), AUPRC {_fmt(auprc)} (prevalencia 4,5 %) y localiza la oración errónea "
            f"en {_pct(loc)} ({aciertos}/{n_notas} notas). Se entregaron comparación tripartita, índice "
            "FAISS+RAG, demo Streamlit, API FastAPI, frontend React, fallback TF-IDF, evaluaciones por "
            f"tipo de error y sesgo EN-ES. Como avance hacia CITIMED, se implementó anonimización local "
            f"(s10/anonimizador/) y {citimed_eval_txt}; {piloto_infer_txt}. "
            "Corpus clínico anotado y aval ético pendientes. "
            f"Repositorio: {REPO_URL}.",
        )

    # --- §2.2 Herramientas ---
    s22 = _find_paragraph(doc, "El pipeline se implementó") or _find_paragraph(doc, "El pipeline se construyó")
    if s22:
        _replace_paragraph_text(
            s22,
            "El pipeline se implementó con herramientas de código abierto: scikit-learn (TF-IDF + "
            "Regresión Logística), cliente OpenAI-compatible para LLM (s7/llm_client.py), "
            "sentence-transformers + FAISS para RAG (s7/rag_index.py, s7/knowledge/), inferencia "
            "unificada (s7/inferencia.py), demo Streamlit (demo/), API REST FastAPI (api/) y "
            "frontend React (frontend/). Para CITIMED se añadió el servicio de anonimización "
            "(s10/anonimizador/ANONIMIZADOR/: reglas Ecuador + spaCy NER), utilidades compartidas "
            "s10/citimed_utils.py y s10/probar_citimed.py (anonimizar → inferencia ES → eval "
            "cross-domain sobre plantilla etiquetada). Mock LLM activo por defecto; soporte "
            "Ollama on-premise (OPENAI_BASE_URL=http://localhost:11434/v1) para PHI. Trabajo futuro "
            "(no entregado en S10): orquestación LangChain/LlamaIndex, Docker Compose, DVC, MLflow "
            "operativo (requirements-optional.txt, entorno separado). Reproducibilidad: SEED=42, "
            "particiones oficiales MEDEC, métricas en salidas_ajuste/ y salidas_s7/.",
        )

    # --- Tabla métricas §3.2 (Table 1) ---
    tbl_metrics = _find_table_by_header(doc, "Métrica (test MEDEC)")
    if tbl_metrics:
        # Actualizar filas existentes
        for row in tbl_metrics.rows[1:]:
            metric = row.cells[0].text.strip()
            if metric == "ROC-AUC":
                row.cells[2].text = _fmt(roc)
            elif "ROC-AUC (IC" in metric:
                row.cells[2].text = f"{_fmt(roc_ci[1])} – {_fmt(roc_ci[2])}"
            elif "CV 5-fold" in metric:
                row.cells[2].text = f"{_fmt(cv_mean)} ± {_fmt(cv_sd)}"
            elif "Localización" in metric:
                row.cells[2].text = _fmt(loc)
                row.cells[3].text = f"Acierta la oración en {aciertos}/{n_notas} notas"
        # Añadir fila AUPRC si no existe
        has_auprc = any("AUPRC" in r.cells[0].text for r in tbl_metrics.rows)
        if not has_auprc:
            _add_row(
                tbl_metrics,
                [
                    "AUPRC (oración)",
                    "—",
                    _fmt(auprc),
                    f"IC 95 % {_fmt(auprc_ci[1])}–{_fmt(auprc_ci[2])}; métrica clave con 4,5 % positivos",
                ],
            )

    # --- Placeholder Table 0 (CITIMED / tripartita / MLflow) ---
    for tbl in doc.tables:
        cell = tbl.rows[0].cells[0].text
        if "Sustituir/complementar MEDEC" in cell or "Corpus CITIMED en español: AVANCE PILOTO" in cell:
            tbl.rows[0].cells[0].text = (
                "Actualización S10:\n"
                "• Comparación tripartita TF-IDF / LLM / LLM+RAG: IMPLEMENTADA "
                f"(s7/eval_tripartita.py; subset n={n_trip}, mock LLM={mock}).\n"
                f"  TF-IDF ROC-AUC={_fmt(tf_t.get('roc_auc', roc))}; LLM zero-shot={_fmt(llm_t.get('roc_auc', 0.5))}; "
                f"LLM+RAG={_fmt(rag_t.get('roc_auc', 0.5))}. McNemar TF-IDF vs LLM: p={_fmt(mcnemar_p)}.\n"
                "• Corpus CITIMED en español: AVANCE PILOTO S10 — anonimizador funcional "
                f"(s10/anonimizador/); {piloto_infer_txt}; {citimed_eval_txt}. "
                "Corpus clínico anotado y aval comité de ética: PENDIENTE.\n"
                "• MLflow operativo: PLAN FUTURO (requirements-optional.txt; venv separado numpy<2)."
            )
        elif "Reportar resultados de la comparación tripartita" in cell:
            filas_tipo = m["por_tipo"].get("tipos", [])
            tipo_txt = "; ".join(
                f"{t['error_type']}: recall={_fmt(t['recall'])}, loc={_fmt(t['localizacion_top1'])}"
                for t in filas_tipo[:5]
            )
            tbl.rows[0].cells[0].text = (
                "Comparación tripartita (test MEDEC, s7/eval_tripartita.py):\n"
                f"• TF-IDF: ROC-AUC={_fmt(tf_t.get('roc_auc', roc))}, AUPRC={_fmt(tf_t.get('auprc', auprc))}, "
                f"latencia≈{trip.get('tfidf', {}).get('latency_ms_per_oracion', 0.4):.2f} ms/oración.\n"
                f"• LLM zero-shot ({'mock' if mock else 'API'}): ROC-AUC={_fmt(llm_t.get('roc_auc', 0.5))} "
                "(sin discriminación en mock; evaluación con API/Ollama pendiente).\n"
                f"• LLM+RAG: ROC-AUC={_fmt(rag_t.get('roc_auc', 0.5))}.\n"
                f"Análisis por ErrorType (TF-IDF): {tipo_txt}.\n"
                f"Sesgo EN-ES (subset n={idioma.get('subset_n', 200)}): ΔAUC={_fmt(delta_auc)}. "
                f"TF-IDF stop_words español: test AUC={_fmt(tfidf_es.get('roc_auc', roc))}."
            )
        elif "Incorporar explicaciones locales" in cell:
            tbl.rows[0].cells[0].text = (
                "Interpretabilidad entregada (S7-S10):\n"
                "• SHAP/LIME sobre TF-IDF: s7/interpretabilidad.py → salidas_s7/shap_*.png, lime_*.png "
                "(ver s7/docs/reporte_interpretabilidad_etica.md).\n"
                "• Demo Streamlit y API FastAPI muestran: oración top-1, scores por brazo, contexto RAG "
                "recuperado (demo/components/results_panel.py).\n"
                "• Formato de reporte al médico: oración señalada + tipo de error + scores + evidencia RAG; "
                "corrección sugerida automática: plan futuro."
            )
        elif "Formalizar el protocolo de anonimización" in cell:
            tbl.rows[0].cells[0].text = (
                "Riesgos y mitigaciones implementadas:\n"
                "• Servicio de anonimización local (s10/anonimizador/): reglas + NER, pseudonimización "
                "consistente, fechas desplazadas, auditoría con hashes (sin PHI en claro). Revisión humana "
                "por muestreo recomendada antes de liberar datos.\n"
                "• Mock LLM por defecto; consentimiento PHI en demo (demo/components/security.py).\n"
                "• .gitignore: CSV/salidas CITIMED generadas, historia_ejemplo.txt local; solo "
                "ejemplos sintéticos versionados (ejemplos/historia_ejemplo.sintetico.txt).\n"
                "• Pipeline aborta si PHI residual (s10/citimed_utils.py); tests en "
                "s10/test_citimed_pipeline.py.\n"
                "• Auditoría SHA-256 sin texto clínico (salidas_s7/audit.log).\n"
                "• Fallback TF-IDF si falla API LLM (s7/test_fallback.py; modo_degradado=True).\n"
                "• Evaluación sesgo EN-ES: s7/eval_idioma.py, s7/eval_tfidf_idioma.py.\n"
                "• Protocolo formal CITIMED + aprobación comité de ética: pendiente organizacional."
            )

    # --- Tabla trabajo restante §7 (Table 5) ---
    tbl_work = _find_table_by_header(doc, "#")
    if tbl_work:
        status_map = {
            "Comparación tripartita": "Implementado (S7/S10)",
            "Índice FAISS": "Implementado (s7/rag_index.py)",
            "MLflow operativo": "Plan futuro",
            "Corpus de CITIMED": "Piloto S10 (anonimizador + 1 historia)",
            "Interpretabilidad con cita": "Implementado (SHAP/LIME + RAG demo)",
            "Redacción final": "S10 consolidado",
        }
        for row in tbl_work.rows[1:]:
            tarea = row.cells[1].text.strip()
            for key, estado in status_map.items():
                if key.lower() in tarea.lower():
                    row.cells[2].text = estado
                    row.cells[3].text = "S10"
                    break

    # --- §5 Interpretabilidad ---
    s5a = _find_paragraph(doc, "La interpretabilidad es un requisito")
    if s5a:
        _replace_paragraph_text(
            s5a,
            "La interpretabilidad es un requisito, no un adorno: un asistente clínico debe justificar "
            "cada alerta. El proyecto la aborda en tres niveles. Primero, coeficientes TF-IDF auditable "
            "en Regresión Logística. Segundo, localización top-1 del "
            f"{_pct(loc)} (oración concreta verificable). Tercero, explicaciones SHAP/LIME "
            "(s7/interpretabilidad.py) y contexto RAG citado en la demo y la API.",
        )
    s5b = _find_paragraph(doc, "Para la fase LLM + RAG, la interpretabilidad")
    if s5b:
        _replace_paragraph_text(
            s5b,
            "En LLM+RAG, la demo muestra fragmentos recuperados de guías clínicas (s7/knowledge/) "
            "junto a cada score. Faithfulness RAG formal y corrección sugerida automática quedan "
            "como trabajo futuro; las figuras SHAP/LIME están en s10/evidencias/capturas/.",
        )

    # --- §7 Conclusiones ---
    s7a = _find_paragraph(doc, "Hasta la semana 10") or _find_paragraph(doc, "Hasta la semana 8")
    if s7a:
        _replace_paragraph_text(
            s7a,
            "Hasta la semana 10, el proyecto cuenta con un problema bien delimitado, arquitectura "
            "implementada y evaluada sobre MEDEC, pipeline reproducible (s6/, s7/, api/, demo/) y "
            f"línea base cuantitativa: ROC-AUC {_fmt(roc)}, AUPRC {_fmt(auprc)}, localización "
            f"{_pct(loc)}. La comparación tripartita confirma ventaja de TF-IDF en modo mock "
            f"(LLM AUC≈0,5); evaluación con API/Ollama real pendiente. Como avance CITIMED (no corpus "
            f"final), se validó anonimización local, {piloto_infer_txt} y {citimed_eval_txt}. "
            "API FastAPI verificada (GET /health, POST /generar con alerta en ejemplo medicación). "
            "El criterio de éxito para piloto clínico: corpus anotado "
            "con aval ético, mantener recall en farmacoterapia/diagnóstico, desplegar con Ollama "
            "on-premise y fallback transparente.",
        )
    s7b = _find_paragraph(doc, "Trabajo restante para la versión final:")
    if s7b:
        _replace_paragraph_text(
            s7b,
            "Trabajo futuro (post-S10): corpus CITIMED clínico anotado con aval ético, ampliar lote "
            "anonimizado, eval tripartita con API/Ollama real, cascada TF-IDF→LLM codificada, "
            "Docker/DVC/MLflow, corrección sugerida y faithfulness RAG.",
        )

    # --- Anexo técnico ---
    anexo = _find_paragraph(doc, "Todos los resultados son reproducibles")
    if anexo:
        _replace_paragraph_text(
            anexo,
            "Todos los resultados son reproducibles desde la raíz del repositorio MIA_Deteccion_HC/ "
            f"(SEED=42, particiones oficiales MEDEC). Estructura: s6/ (TF-IDF), s7/ (LLM, RAG, evals), "
            "api/ (FastAPI), frontend/ (React), demo/ (Streamlit), s10/ (evidencias, informe, anonimizador). "
            "Salidas: salidas_ajuste/, salidas_s7/ (incl. prueba_citimed.json). Comandos: "
            "python s6/modelo_ajustado.py; python s7/eval_tripartita.py --mock-llm --max-oraciones 500; "
            "python s7/test_fallback.py; python s10/probar_citimed.py; "
            "python -m pytest s10/test_citimed_pipeline.py -q; streamlit run demo/app.py; "
            "uvicorn api.main:app --port 8000 (verificar GET /health y POST /generar en /docs). "
            f"Evidencias: s10/evidencias/. Repositorio: {REPO_URL}. "
            "Nota entrega: commit final S10 (anonimizador + cleanup CITIMED + informe) pendiente en git.",
        )

    # --- §3.3 Piloto CITIMED (avance, no corpus final) ---
    ultimo_piloto = None
    if piloto["tiene_piloto"] and not _find_paragraph(doc, "3.3. Piloto CITIMED"):
        eval_párrafo = (
            f"Eval cross-domain sobre plantilla odontológica etiquetada "
            f"(n={piloto['n_oraciones_eval']} oraciones): ROC-AUC = {_fmt_opt(piloto['roc'])}, "
            f"AUPRC = {_fmt_opt(piloto['auprc'])}."
            if piloto["tiene_eval"]
            else "Eval cross-domain sobre plantilla odontológica: pendiente de ejecutar s10/probar_citimed.py."
        )
        ultimo_piloto = _insert_paragraphs_after(
            doc,
            "Un matiz importante: al reagregar",
            [
                ("Heading 2", "3.3. Piloto CITIMED — anonimización e integración (avance S10)"),
                (
                    "Normal",
                    "Como avance hacia el despliegue en CITIMED (no como corpus clínico final), se "
                    "implementó el servicio de anonimización en s10/anonimizador/ANONIMIZADOR/ "
                    "(reglas para identificadores ecuatorianos + NER spaCy, pseudonimización consistente, "
                    "fechas desplazadas y auditoría sin PHI en claro). El script s10/probar_citimed.py "
                    "ejecuta el flujo: historia → anonimizar → inferencia en español (TF-IDF + LLM mock). "
                    "La evaluación cross-domain usa solo la plantilla odontológica etiquetada; las historias "
                    "anonimizadas del piloto se reportan aparte (sin mezclar labels artificiales).",
                ),
                (
                    "Normal",
                    f"Inferencia piloto: {piloto['n_historias']} historia(s) anonimizada(s), "
                    f"{piloto['n_oraciones_inferencia']} oraciones segmentadas. {eval_párrafo} "
                    "Estas cifras son preliminares (subset pequeño) y no sustituyen validación con corpus "
                    "CITIMED anotado y aval del comité de ética. Reporte: salidas_s7/prueba_citimed.json.",
                ),
                (
                    "Normal",
                    "Limitaciones del piloto: historias de ejemplo sintéticas; métricas cross-domain "
                    "sensibles al tamaño muestral; anonimización automática requiere revisión humana por "
                    "muestreo antes de uso clínico. Trabajo pendiente: lote anonimizado institucional, "
                    "anotación de inconsistencias y evaluación con Ollama on-premise.",
                ),
            ],
        )

    # --- §3.4 Verificación FastAPI y calidad código ---
    if not _find_paragraph(doc, "3.4. Verificación FastAPI"):
        ancla_34 = "Limitaciones del piloto"
        if _find_paragraph(doc, ancla_34) is None:
            ancla_34 = "Un matiz importante: al reagregar"
        _insert_paragraphs_after(doc, ancla_34, _bloque_verificacion_s10())

    # --- Figuras embebidas ---
    for anchor, path, caption in [
        ("Figura 1 (S6)", EVID / "figura_ajuste.png", "Figura 1. Curvas ROC/PR y lift — baseline vs ajustado (S6)."),
        ("Figura 2 (S7)", CAP / "figura_comparacion_tripartita.png", "Figura 2. Comparación tripartita TF-IDF / LLM / LLM+RAG."),
    ]:
        _insert_picture_after_simple(doc, anchor, path, caption)

    if ultimo_piloto is not None:
        _insert_picture_after_paragraph(
            ultimo_piloto,
            CAP / "heatmap_recall_por_tipo.png",
            "Figura 3. Recall por ErrorType (TF-IDF).",
        )
    else:
        _insert_picture_after_simple(
            doc,
            "Un matiz importante: al reagregar",
            CAP / "heatmap_recall_por_tipo.png",
            "Figura 3. Recall por ErrorType (TF-IDF).",
        )

    _insert_picture_after_simple(
        doc,
        "En LLM+RAG, la demo muestra",
        CAP / "shap_summary_bar.png",
        "Figura 4. SHAP summary (TF-IDF).",
    )

    anexo_para = _find_paragraph(doc, "Evidencias: s10/evidencias/")
    if anexo_para is None:
        anexo_para = _find_paragraph(doc, "salidas_s7/")
    if anexo_para:
        for path, cap in [
            (CAP / "demo_analisis.png", "Figura 5. Demo — análisis ejemplo medicación."),
            (CAP / "demo_metricas.png", "Figura 6. Demo — métricas ROC-AUC / AUPRC."),
            (CAP / "demo_fallback.png", "Figura 7. Modo degradado / fallback TF-IDF."),
        ]:
            _insert_picture_after_simple(doc, "Evidencias: s10/evidencias/", path, cap)

    doc.save(str(OUT_DOCX))
    print(f"[ok] Word (formato S8 actualizado) -> {OUT_DOCX}")


def export_pdf() -> None:
    """Exporta PDF preservando formato Word (LibreOffice, MS Word COM o docx2pdf)."""
    # 1) MS Word COM (Windows)
    try:
        import comtypes.client  # type: ignore

        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(OUT_DOCX.resolve()))
        doc.SaveAs(str(OUT_PDF.resolve()), FileFormat=17)  # wdFormatPDF
        doc.Close()
        word.Quit()
        print(f"[ok] PDF (MS Word) -> {OUT_PDF}")
        return
    except Exception as e:
        print(f"[info] Word COM no disponible: {e}")

    # 2) LibreOffice headless
    for cmd in (
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(DOCS), str(OUT_DOCX)],
        [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(DOCS),
            str(OUT_DOCX),
        ],
    ):
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            pdf_lo = DOCS / (OUT_DOCX.stem + ".pdf")
            if proc.returncode == 0 and pdf_lo.exists():
                if pdf_lo != OUT_PDF:
                    shutil.move(str(pdf_lo), str(OUT_PDF))
                print(f"[ok] PDF (LibreOffice) -> {OUT_PDF}")
                return
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # 3) docx2pdf
    try:
        import docx2pdf  # type: ignore

        docx2pdf.convert(str(OUT_DOCX), str(OUT_PDF))
        print(f"[ok] PDF (docx2pdf) -> {OUT_PDF}")
        return
    except Exception:
        pass

    print(
        "[warn] No se pudo exportar PDF automático. "
        f"Abra {OUT_DOCX} en Word → Guardar como PDF para Moodle."
    )


def main() -> None:
    # Asegurar capturas demo
    if not (CAP / "demo_analisis.png").exists():
        print("[info] Generando capturas demo...")
        from s10.capture_demo import main as cap_main

        cap_main()

    metrics = load_metrics()
    if not metrics["tfidf"]:
        print("[warn] metricas_ajuste.json no encontrado; ejecutar python s6/modelo_ajustado.py")

    update_from_s8(metrics)
    export_pdf()


if __name__ == "__main__":
    main()
