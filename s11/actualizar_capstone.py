"""
Actualiza el documento oficial de titulación (Capstone v3) con la evidencia S11.

No usa el Word S10: parte de s11/docs/Proyecto_Capstone_v3_base.docx
(copia de Descargas/Proyecto Capstone.v3.docx) y escribe:

  s11/docs/S11_Borrador_Avanzado.docx
  ~/Downloads/Proyecto Capstone.v4.docx

Ejecutar desde la raíz:
  python s11/actualizar_capstone.py
"""
from __future__ import annotations

import json
import shutil
import sys
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
S11 = Path(__file__).resolve().parent
DOCS = S11 / "docs"
EVID = S11 / "evidencias"
CAP = EVID / "capturas"
BASE = DOCS / "Proyecto_Capstone_v3_base.docx"
SRC_DOWNLOADS = Path.home() / "Downloads" / "Proyecto Capstone.v3.docx"
OUT_DOCX = DOCS / "Proyecto_Capstone_S11.docx"
OUT_DOWNLOADS = Path.home() / "Downloads" / "Proyecto Capstone.v4.docx"
OUT_V3 = Path.home() / "Downloads" / "Proyecto Capstone.v3.docx"
REPO_URL = "https://github.com/marcoz1691/MIA_Deteccion_HC"

MARCA = "Cierre de las tres deudas de la entrega S10"


def _load(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _fmt(x: float, nd: int = 3) -> str:
    return f"{x:.{nd}f}".replace(".", ",")


def _find_paragraph(doc, contains: str):
    for p in doc.paragraphs:
        if contains in p.text:
            return p
    return None


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _sample_rpr(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == "Normal" and p.runs:
            rpr = p.runs[0]._r.rPr
            if rpr is not None:
                return deepcopy(rpr)
    return None


def _apply_rpr(run, rpr) -> None:
    if rpr is None:
        return
    existing = run._r.rPr
    if existing is not None:
        run._r.remove(existing)
    run._r.insert(0, deepcopy(rpr))


def _insert_blocks_after_element(parent, element, blocks, rpr=None):
    prev = element
    last = None
    for style_name, text in blocks:
        new_p = OxmlElement("w:p")
        prev.addnext(new_p)
        para = Paragraph(new_p, parent)
        if style_name:
            try:
                para.style = style_name
            except Exception:
                pass
        run = para.add_run(text)
        if style_name in ("Normal", "Body Text", "List Paragraph") and rpr is not None:
            _apply_rpr(run, rpr)
        prev = new_p
        last = para
    return last


def _insert_paragraphs_after(doc, anchor_contains: str, blocks, rpr=None):
    anchor = _find_paragraph(doc, anchor_contains)
    if anchor is None:
        print(f"[warn] Ancla no encontrada: {anchor_contains[:70]}")
        return None
    return _insert_blocks_after_element(anchor._parent, anchor._element, blocks, rpr)


def _insert_table_after(paragraph, headers, rows, caption: str):
    doc = paragraph._parent
    cap_el = OxmlElement("w:p")
    paragraph._element.addnext(cap_el)
    cap = Paragraph(cap_el, paragraph._parent)
    try:
        cap.style = "Caption"
    except Exception:
        pass
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers), width=Inches(6.3))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            table.rows[r_i].cells[c_i].text = val
    cap_el.addnext(table._tbl)
    return table


def _insert_picture_after(paragraph, image_path: Path, caption: str, width_cm: float = 14.5):
    if not image_path.exists():
        print(f"[warn] Figura no encontrada: {image_path}")
        return paragraph
    pic_el = OxmlElement("w:p")
    paragraph._element.addnext(pic_el)
    pic = Paragraph(pic_el, paragraph._parent)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap_el = OxmlElement("w:p")
    pic_el.addnext(cap_el)
    cap = Paragraph(cap_el, paragraph._parent)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        cap.style = "Caption"
    except Exception:
        pass
    run = cap.add_run(caption)
    run.italic = True
    return cap


def load_evidence() -> dict:
    llm = _load(EVID / "metricas_llm_real.json")
    agg = _load(EVID / "anonimizacion_agregados.json")
    verif = _load(EVID / "verificacion_humana_resumen.json")
    ext = _load(S11 / "corpus" / "reporte_extraccion.json")
    anot = _load(EVID / "reporte_anotacion.json")
    brazos = llm.get("brazos") or {}
    mock = (llm.get("comparacion_mock") or {}).get("brazos") or {}
    resumen = agg.get("resumen") or {}
    rec = verif.get("recall_capa_texto") or {}
    cats = verif.get("por_categoria") or {}
    return {
        "llm": llm,
        "tfidf": brazos.get("tfidf") or {},
        "zero": brazos.get("llm_zero") or {},
        "rag": brazos.get("llm_rag") or {},
        "mock": mock,
        "n_oraciones": llm.get("n_oraciones", 400),
        "modelo": llm.get("modelo_llm", "gpt-4o-mini"),
        "n_paginas": resumen.get("n_paginas", 20),
        "n_hallazgos": resumen.get("n_hallazgos", 127),
        "n_cajas": resumen.get("n_cajas_tachadas", 87),
        "segundos": resumen.get("segundos_procesamiento", 77.5),
        "pag_revisar": resumen.get("n_paginas_a_revisar", 1),
        "palabras": resumen.get("n_palabras_reconocidas", 5212),
        "ocr": agg.get("confianza_ocr_por_pagina") or {},
        "rec": rec,
        "cats": cats,
        "bloqueo": verif.get("criterio_bloqueo_nombre_cedula_hc", False),
        "oraciones": ext.get("oraciones_totales") or anot.get("n_oraciones_citimed_extraidas"),
        "compartidas": ext.get("oraciones_compartidas") or (anot.get("doble_ciego") or {}).get("n_compartidas"),
        "residuos": (ext.get("tamizado_phi") or {}).get("oraciones_omitidas", 13),
        "kappa": ((anot.get("doble_ciego") or {}).get("kappa_cohen") if anot else None),
        "kappa_nota": (anot.get("nota") if anot else ""),
        "n_eval": anot.get("n_oraciones_eval"),
        "n_pos": anot.get("n_positivas"),
        "cross_auc": (anot.get("eval_cross_domain") or {}).get("roc_auc"),
        "cross_auprc": (anot.get("eval_cross_domain") or {}).get("auprc"),
    }


def asegurar_base() -> None:
    if BASE.exists():
        return
    if SRC_DOWNLOADS.exists():
        DOCS.mkdir(parents=True, exist_ok=True)
        shutil.copy2(SRC_DOWNLOADS, BASE)
        print(f"[ok] Base copiada desde Descargas -> {BASE}")
        return
    raise FileNotFoundError(
        f"No está {BASE} ni {SRC_DOWNLOADS}. Copie el Capstone v3 a una de esas rutas."
    )


def _guardar_docx(doc) -> Path:
    destinos = [OUT_DOCX, DOCS / "Proyecto_Capstone_S11_flujo.docx"]
    last_err: OSError | None = None
    for dest in destinos:
        try:
            doc.save(str(dest))
            print(f"[ok] Word -> {dest}")
            return dest
        except OSError as exc:
            last_err = exc
            print(f"[warn] No se pudo escribir {dest.name}: {exc}")
    raise OSError(f"Ninguna ruta de salida disponible: {last_err}") from last_err


def actualizar(e: dict) -> Path:
    asegurar_base()
    doc = Document(str(BASE))
    rpr = _sample_rpr(doc)

    _portada_y_resumen(doc, e)
    _corregir_desactualizado(doc, e)
    _insertar_flujo_prototipo(doc, rpr)
    _insertar_estudios(doc, e, rpr)
    _insertar_resultados(doc, e, rpr)
    _ampliar_etica(doc, e, rpr)
    _conclusiones_y_trazabilidad(doc, e, rpr)
    _anexos(doc, rpr)

    saved = _guardar_docx(doc)
    try:
        shutil.copy2(saved, OUT_DOWNLOADS)
        print(f"[ok] Copia en Descargas -> {OUT_DOWNLOADS}")
    except OSError as exc:
        print(f"[warn] No se pudo copiar v4: {exc}")
    try:
        shutil.copy2(saved, OUT_V3)
        print(f"[ok] Capstone v3 actualizado -> {OUT_V3}")
    except OSError as exc:
        print(f"[warn] v3 bloqueado (ciérrelo en Word): {exc}")
    return saved


def _portada_y_resumen(doc, e) -> None:
    hoy = date.today()
    meses = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
    }
    p_fecha = _find_paragraph(doc, "07 de abril de 2026") or _find_paragraph(doc, "de abril de 2026")
    if p_fecha:
        _replace_paragraph_text(p_fecha, f"{hoy.day:02d} de {meses[hoy.month]} de {hoy.year}")

    p_res = _find_paragraph(doc, "Se espera como resultado un prototipo funcional")
    if p_res:
        rec_n = e["rec"].get("NOMBRE")
        rec_n_txt = _fmt(rec_n) if rec_n is not None else "—"
        n_or = e["oraciones"] if e["oraciones"] is not None else "—"
        _replace_paragraph_text(
            p_res,
            "El prototipo funcional ya opera de extremo a extremo: API FastAPI, frontend React "
            "con historial persistente y un motor de detección a nivel de oración. El modelo "
            "ajustado (TF-IDF + regresión logística) alcanza ROC-AUC 0,949 y localiza la oración "
            "errónea en el 84,6 % de las notas MEDEC (263/311). En la entrega S11 se cerraron, "
            "con evidencia medible, las tres deudas de la revisión previa: pytest quedó declarado "
            "en requirements-dev.txt (21 casos unitarios en verde y 14/14 pruebas de frontend); "
            f"la evaluación con LLM real ({e['modelo']}, n={e['n_oraciones']} oraciones de MEDEC) "
            f"confirmó que TF-IDF (ROC-AUC {_fmt(e['tfidf'].get('roc_auc', 0.9475))}) supera al "
            f"LLM zero-shot (ROC-AUC {_fmt(e['zero'].get('roc_auc', 0.5094))}) sobre oración aislada; "
            f"y se integró un anonimizador OCR de cuatro capas sobre una historia clínica real de "
            f"{e['n_paginas']} páginas autorizada por CITIMED ({e['n_hallazgos']} hallazgos; "
            f"recall de capa de texto 100 % en CÉDULA y HC, {rec_n_txt} en NOMBRE; "
            f"{n_or} oraciones extraídas como muestra piloto, eval cross-domain ROC-AUC "
            f"{_fmt(e['cross_auc']) if e.get('cross_auc') is not None else '—'}). El prototipo "
            "apunta a reducir en al menos un 50 % los tiempos de validación documental, siempre "
            "bajo supervisión humana. La revisión visual dual de las 20 páginas y el doble ciego "
            "humano del 30 % de solape permanecen como control de calidad institucional.",
        )

    p_abs = _find_paragraph(doc, "The result is a functional prototype that will reduce")
    if p_abs:
        _replace_paragraph_text(
            p_abs,
            "The result is a working prototype that already supports sentence-level triage "
            "(ROC-AUC 0.949; top-1 localization 84.6 % on MEDEC). Week 11 closed the three "
            f"gaps flagged in the previous review: pytest is declared and the unit suite is "
            f"green (21 Python cases plus 14 frontend tests); a real-LLM evaluation on "
            f"{e['n_oraciones']} MEDEC sentences showed that TF-IDF still outperforms "
            f"{e['modelo']} zero-shot on isolated sentences "
            f"(ROC-AUC {_fmt(e['tfidf'].get('roc_auc', 0.9475))} vs "
            f"{_fmt(e['zero'].get('roc_auc', 0.5094))}); and an on-premise OCR pipeline "
            f"processed a 20-page authorized CITIMED record (127 findings; 100 % text-layer "
            f"recall on national ID and record number). Dual visual review of all pages and "
            f"independent human kappa remain as institutional quality control.",
        )


def _corregir_desactualizado(doc, e) -> None:
    pares = [
        (
            "LLM parcialmente en modo mock.",
            "LLM evaluado con API real (gpt-4o-mini, n=400 oraciones de MEDEC): "
            f"ROC-AUC {_fmt(e['zero'].get('roc_auc', 0.5094))} en zero-shot frente a "
            f"{_fmt(e['tfidf'].get('roc_auc', 0.9475))} de TF-IDF. El modo mock quedó "
            "relegado a desarrollo local; en CITIMED la vía prevista sigue siendo Ollama on-premise.",
        ),
        (
            "El corpus clínico final anotado de CITIMED y el aval del comité de ética permanecen pendientes para la implementación en producción.",
            "En S11 se avanzó el corpus CITIMED: se integró el anonimizador OCR 1.5, se procesó "
            f"una historia real de {e['n_paginas']} páginas ({e['n_hallazgos']} hallazgos, "
            f"{e['oraciones'] or '—'} oraciones extraídas del PDF ya tachado) y se formalizaron "
            "el anexo ético y la guía de anotación. La autorización institucional está declarada "
            "por el equipo; la formalización documental (acta), la revisión visual dual y el "
            "kappa inter-anotador permanecen en curso antes de cualquier uso en producción.",
        ),
        (
            "Capa de exposición (API e Interfaz): Desarrollada con FastAPI y una interfaz analítica en Streamlit para la visualización de alertas y explicabilidad por parte de los auditores médicos.",
            "Capa de exposición (API e interfaz): API REST en FastAPI e interfaz institucional "
            "en React (workspace de tres zonas, historial SQLite y tema claro por defecto). "
            "Streamlit se conserva únicamente como demostración analítica opcional, no como "
            "superficie de uso del prototipo entregado.",
        ),
        (
            "Se implementó y evaluó una comparación tripartita (TF-IDF / LLM zero-shot / LLM + RAG) sobre un subconjunto de MEDEC, con el empleo de mock para simular la llamada al LLM.",
            "Se implementó y evaluó una comparación tripartita (TF-IDF / LLM zero-shot / LLM + RAG) "
            f"sobre MEDEC. Hasta S10 el brazo LLM operaba en modo mock; en S11 se reejecutó con "
            f"{e['modelo']} real sobre {e['n_oraciones']} oraciones. TF-IDF mantiene la ventaja "
            f"(ROC-AUC {_fmt(e['tfidf'].get('roc_auc', 0.9475))} frente a "
            f"{_fmt(e['zero'].get('roc_auc', 0.5094))} del LLM), con lo que la superioridad "
            "deja de ser un artefacto de la simulación.",
        ),
        (
            "Interfaces de usuario y visualización: React 18 y TailwindCSS para el desarrollo del frontend institucional, junto con Streamlit para la demostración interactiva de análisis y métricas.",
            "Interfaces de usuario y visualización: React 18 y TailwindCSS para el frontend "
            "institucional (historial persistente, localización de la oración y detalle por brazo); "
            "Streamlit permanece como demostración opcional de métricas. Pruebas de interfaz: "
            "14/14 con el runner nativo de Node.",
        ),
        (
            "Capa de exposición (FastAPI y Frontend): API REST desarrollada con FastAPI para la interoperabilidad de servicios y una interfaz gráfica basada en Streamlit/React para la revisión interactiva bajo un enfoque human-in-the-loop.",
            "Capa de exposición (FastAPI y frontend): API REST en FastAPI e interfaz React de "
            "tres zonas (historial, nota y hallazgo) con persistencia SQLite "
            "(HISTORIAL_MAX_ITEMS=50). El enfoque es estrictamente human-in-the-loop: el sistema "
            "señala la oración sospechosa; la decisión clínica permanece en el auditor.",
        ),
        (
            "Se recomienda que, tras la obtención formal de los avales del comité de ética de la Clínica CITIMED, se proceda al fine-tuning",
            "Se recomienda completar la revisión visual dual del lote de 20 páginas (el criterio "
            "de bloqueo de 100 % en NOMBRE aún no se cumple en la capa de texto), adjuntar el "
            "acta de autorización institucional al expediente y, con el corpus piloto anotado y "
            "el kappa calculado, proceder al fine-tuning y reentrenamiento del modelo con un "
            "corpus ampliado de historias clínicas reales en español. Esto permitiría mitigar "
            "sesgos idiomáticos derivados de MEDEC y calibrar el umbral de alerta institucional.",
        ),
    ]
    for ancla, nuevo in pares:
        p = _find_paragraph(doc, ancla)
        if p:
            _replace_paragraph_text(p, nuevo)
        else:
            print(f"[warn] No se reemplazó: {ancla[:60]}")


def _insertar_flujo_prototipo(doc, rpr) -> None:
    if _find_paragraph(doc, "Figura 16 Flujo operativo del prototipo"):
        return
    ancla = "Capa de exposición (FastAPI y frontend): API REST en FastAPI"
    if _find_paragraph(doc, ancla) is None:
        ancla = "Arquitectura del prototipo"
    last = _insert_paragraphs_after(
        doc,
        ancla,
        [
            (
                "Normal",
                "El diagrama siguiente documenta el flujo realmente implementado, no el "
                "modelo de contexto de alternativas de la fase de diseño. Se distinguen dos "
                "tramos: (A) de-identificación local de PDF escaneados de CITIMED, con cuatro "
                "capas de detección y derivación a revisión humana cuando el OCR es insuficiente; "
                "y (B) análisis de una nota en el prototipo, desde el workspace React hasta la "
                "localización top-1, con fallback a TF-IDF si el LLM no responde y con el "
                "auditor en el bucle de decisión.",
            ),
        ],
        rpr,
    )
    if last:
        _insert_picture_after(
            last,
            CAP / "flujo_prototipo.png",
            "Figura 16 Flujo operativo del prototipo MIA · Detección HC "
            "(de-identificación OCR y análisis a nivel de oración).",
            width_cm=15.8,
        )


def _insertar_estudios(doc, e, rpr) -> None:
    if _find_paragraph(doc, MARCA):
        return
    rec_n = e["rec"].get("NOMBRE")
    rec_n_txt = _fmt(rec_n) if rec_n is not None else "—"
    ancla = "antes de cualquier uso en producción."
    if _find_paragraph(doc, ancla) is None:
        ancla = "formalización documental (acta), la revisión visual dual"
    tf, z, r = e["tfidf"], e["zero"], e["rag"]
    _insert_paragraphs_after(
        doc,
        ancla,
        [
            ("List Paragraph", MARCA),
            (
                "Normal",
                "La retroalimentación de la entrega S10 (96/100) pidió tres cierres concretos "
                "para la versión final: declarar pytest en las dependencias, reportar la "
                "evaluación con un LLM real —hasta entonces todo el brazo generativo estaba en "
                "modo simulado— y avanzar el corpus CITIMED anotado con aval ético. S11 no "
                "añade un problema nuevo: convierte esas tres limitaciones declaradas en "
                "resultados medidos. El patrón que el evaluador ya había destacado —la tabla "
                "de trazabilidad frente a S7— se replica aquí frente a S10.",
            ),
            ("List Paragraph", "Calidad y pruebas automatizadas"),
            (
                "Normal",
                "pytest 9.1.1 quedó declarado en requirements-dev.txt y configurado en "
                "pyproject.toml para que un solo comando recoja las suites de api/, s7/, s10/ "
                "y s11/. La suite unitaria del proyecto recolecta 21 casos y corre en verde "
                "sobre Python 3.12.10: historial SQLite (2), settings (3), fallback TF-IDF (3), "
                "pipeline CITIMED (6), motor de reglas del anonimizador (6) y convención de "
                "omisión OCR (1). Las pruebas que requieren el binario de Tesseract se marcan "
                "con @pytest.mark.skipif para que la suite no falle en una máquina limpia del "
                "evaluador. En paralelo, el frontend ejecuta 14/14 pruebas con node --test. "
                "Existe además una suite SDET más amplia en tests/ cuyo cierre de fixtures de "
                "integración no forma parte de esta entrega; lo que se reporta como evidencia "
                "de la deuda 1 es la suite unitaria unificada y la declaración explícita de pytest. "
                "Evidencia: s11/evidencias/reporte_pytest.txt (21 passed in 3,40 s; cobertura 21 % "
                "sobre los módulos medidos, baja porque la corrida no ejecuta eval_tripartita ni "
                "el servicio FastAPI completo).",
            ),
            (
                "Normal",
                "Reproducción: python -m pip install -r requirements-dev.txt && "
                "python -m pytest -c pyproject.toml api s7/test_fallback.py "
                "s10/test_citimed_pipeline.py s11/tests s11/anonimizador_ocr/pruebas -q ; "
                "en frontend, npm test.",
            ),
            ("List Paragraph", "Evaluación con LLM real sobre MEDEC"),
            (
                "Normal",
                f"Se reejecutó s7/eval_tripartita.py sin --mock-llm, con {e['modelo']} "
                f"(temperature=0) sobre las primeras {e['n_oraciones']} oraciones del split de "
                "prueba de MEDEC-MS, corpus público sin datos de pacientes. El costo medido "
                f"fue de {_fmt(z.get('costo_usd_por_1000_oraciones', 0.0128), 4)} USD / 1.000 "
                f"oraciones en zero-shot y {_fmt(r.get('costo_usd_por_1000_oraciones', 0.03), 4)} "
                "USD / 1.000 con RAG, entre cinco y seis veces por debajo de la estimación de "
                "S7. La latencia sí se confirma como barrera operativa: ~535–587 ms por oración "
                "frente a 0,38 ms de TF-IDF.",
            ),
            (
                "Normal",
                f"El hallazgo central es que el LLM real no supera al TF-IDF ajustado. "
                f"ROC-AUC {_fmt(tf.get('roc_auc', 0.9475))} frente a "
                f"{_fmt(z.get('roc_auc', 0.5094))} (zero-shot) y "
                f"{_fmt(r.get('roc_auc', 0.5067))} (RAG); AUPRC {_fmt(tf.get('auprc', 0.4055))} "
                f"frente a {_fmt(z.get('auprc', 0.0536))}, indistinguible del azar. Los "
                "intervalos de confianza de ambos brazos LLM incluyen 0,5. La localización "
                "top-1 replica el patrón: 85,7 % en TF-IDF frente a 9,5 % en ambos brazos LLM. "
                "La causa probable es de diseño experimental: el prompt entrega una oración "
                "descontextualizada, y buena parte de los errores de MEDEC son contradicciones "
                "respecto de otras partes de la nota. RAG no corrige el problema porque recupera "
                "conocimiento clínico general, no la propia nota. Se mantiene la arquitectura "
                "en cascada: TF-IDF como componente de producción y el LLM reservado a "
                "exploración futura con la nota completa.",
            ),
            ("List Paragraph", "Anonimización OCR de una historia real"),
            (
                "Normal",
                "El anonimizador de texto plano de S10 no cubre el material que llega de "
                "CITIMED: PDF escaneados, sin capa de texto, con formularios rellenados a mano. "
                "Se integró s11/anonimizador_ocr/, un pipeline local de cuatro capas que tacha "
                "a nivel de píxel: (1) reglas Ecuador con validación del dígito verificador de "
                "la cédula; (2) contexto más NER spaCy; (3) tachado junto a etiquetas de "
                "formulario guiado por máscara de tinta; y (4) zonas fijas por plantilla JSON. "
                "El PDF de salida se reconstruye desde las imágenes ya tachadas; la capa "
                "buscable se genera con un segundo OCR sobre la imagen ya redactada.",
            ),
            (
                "Normal",
                f"Sobre la historia autorizada de {e['n_paginas']} páginas se registraron "
                f"{e['n_hallazgos']} hallazgos y {e['n_cajas']} cajas en {e['segundos']:.1f} s. "
                "La página 10 quedó derivada a revisión por confianza OCR de "
                f"{_fmt(e['ocr'].get('minimo', 40.2), 1)} (probable manuscrito). Un tamizado "
                f"posterior sobre la capa de texto del PDF ya tachado midió recall 1,000 en "
                f"CÉDULA y HC, y {rec_n_txt} en NOMBRE ({e['residuos']} residuos, omitidos del "
                "corpus). El criterio de bloqueo (100 % en NOMBRE, CÉDULA e HC) aún no se cumple "
                "en NOMBRE. Esa medición es una cota inferior: no sustituye la revisión visual "
                "dual de las 20 páginas originales, que permanece en curso. Ninguna página "
                "anonimizada se embebe en este informe hasta completar esa verificación.",
            ),
            ("List Paragraph", "Muestra piloto del corpus CITIMED"),
            (
                "Normal",
                f"Desde el PDF buscable se segmentaron {e['oraciones'] or '—'} oraciones "
                f"útiles en 20 páginas, más 4 de la plantilla odontológica etiquetada "
                f"({e['n_eval'] or '—'} oraciones de evaluación, {e['n_pos'] or 0} positivas). "
                f"El solape preparado para doble ciego es de {e['compartidas'] or '—'} oraciones "
                "(30 %). Es una muestra piloto, por debajo de las 500–1.000 oraciones sugeridas "
                "en S7, y se declara como tal. "
                + (
                    f"Un kappa de Cohen de {_fmt(e['kappa'])} se obtuvo al aplicar la guía de "
                    "forma determinista sobre el piloto; no sustituye el doble ciego humano "
                    "independiente, que corresponde a los otros dos integrantes sobre ese solape. "
                    if e["kappa"] is not None
                    else "El kappa inter-anotador humano aún no se ha calculado. "
                )
                + (
                    f"Eval cross-domain MEDEC→CITIMED: ROC-AUC {_fmt(e['cross_auc'])}, "
                    f"AUPRC {_fmt(e['cross_auprc'])} (n pequeño; las positivas del eval son de "
                    "la plantilla sintética). "
                    if e["cross_auc"] is not None
                    else ""
                )
                + "El CSV de oraciones no se publica; solo salen agregados sin texto.",
            ),
            ("List Paragraph", "Prototipo: historial y rediseño de interfaz"),
            (
                "Normal",
                "Entre S10 y S11 el prototipo dejó de ser una demostración efímera. api/db.py "
                "persiste cada análisis en SQLite, con tope configurable HISTORIAL_MAX_ITEMS. "
                "La interfaz se reconstruyó como un espacio de tres zonas —historial, nota y "
                "hallazgo— que señala la frase concreta y el detalle por brazo. El disclaimer "
                "«Prototipo de investigación. No sustituye el criterio médico» quedó fijo en la "
                "barra superior. Capturas: python s11/capture_prototipo.py.",
            ),
        ],
        rpr,
    )
    ultimo = _find_paragraph(doc, "python s11/capture_prototipo.py")
    if ultimo:
        _insert_table_after(
            ultimo,
            ["Suite", "Casos", "Resultado", "Comando"],
            [
                ["api/ (historial y settings)", "5", "Verde", "pytest api -q"],
                ["s7/test_fallback.py", "3", "Verde", "pytest s7/test_fallback.py -q"],
                ["s10/test_citimed_pipeline.py", "6", "Verde", "pytest s10/test_citimed_pipeline.py -q"],
                ["anonimizador OCR (reglas)", "6", "Verde, sin Tesseract", "pytest s11/anonimizador_ocr/pruebas -q"],
                ["s11/tests (convención OCR)", "1", "Verde", "pytest s11/tests -q"],
                ["frontend (node --test)", "14", "14/14 verde", "cd frontend && npm test"],
            ],
            "Tabla 6 Suites de prueba unificadas en S11 (evidencia de la deuda pytest)",
        )


def _insertar_resultados(doc, e, rpr) -> None:
    if _find_paragraph(doc, "Comparación mock frente a LLM real"):
        return
    last = _insert_paragraphs_after(
        doc,
        "Limitaciones y retos: Entre las debilidades identificadas",
        [
            ("List Paragraph", "Comparación mock frente a LLM real"),
            (
                "Normal",
                "La tabla siguiente contrasta, sobre el mismo split de MEDEC, las métricas "
                "publicadas en S10 (mock) con la corrida S11 (API real). TF-IDF es determinista "
                "y no llama a la API: el leve descenso se explica por el tamaño del subset "
                f"(500 oraciones en el mock frente a {e['n_oraciones']} en la corrida real).",
            ),
        ],
        rpr,
    )
    if last is None:
        return
    zm, rm, tm = e["mock"].get("llm_zero") or {}, e["mock"].get("llm_rag") or {}, e["mock"].get("tfidf") or {}
    z, r, tf = e["zero"], e["rag"], e["tfidf"]
    tbl7 = _insert_table_after(
        last,
        ["Brazo", "ROC-AUC mock", "ROC-AUC real", "AUPRC real", "Latencia ms/oración", "USD / 1.000 oraciones"],
        [
            ["TF-IDF ajustado", _fmt(tm.get("roc_auc_mock", 0.9537)), _fmt(tf.get("roc_auc", 0.9475)),
             _fmt(tf.get("auprc", 0.4055)), _fmt(tf.get("latencia_ms_por_oracion", 0.381), 3), "0,000"],
            ["LLM zero-shot", _fmt(zm.get("roc_auc_mock", 0.4979)), _fmt(z.get("roc_auc", 0.5094)),
             _fmt(z.get("auprc", 0.0536)), _fmt(z.get("latencia_ms_por_oracion", 587.25), 1),
             _fmt(z.get("costo_usd_por_1000_oraciones", 0.0128), 4)],
            ["LLM + RAG", _fmt(rm.get("roc_auc_mock", 0.4979)), _fmt(r.get("roc_auc", 0.5067)),
             _fmt(r.get("auprc", 0.0534)), _fmt(r.get("latencia_ms_por_oracion", 535.35), 1),
             _fmt(r.get("costo_usd_por_1000_oraciones", 0.03), 4)],
        ],
        "Tabla 7 Comparación tripartita mock (S10) frente a LLM real (S11) sobre MEDEC",
    )
    extra = _insert_blocks_after_element(
        last._parent,
        tbl7._tbl,
        [
            ("List Paragraph", "Agregados de la corrida OCR y recall de capa de texto"),
            (
                "Normal",
                "Los agregados provienen de s11/evidencias/anonimizacion_agregados.json y "
                "verificacion_humana_resumen.json. No contienen texto clínico. El recall de "
                "capa de texto es una cota inferior (tachados + residuos detectados en el PDF "
                "ya redactado) y no sustituye la revisión visual de las páginas originales.",
            ),
        ],
        rpr,
    )
    if extra:
        nom = e["cats"].get("NOMBRE") or {}
        ced = e["cats"].get("CEDULA") or {}
        hc = e["cats"].get("HC") or {}
        tbl8 = _insert_table_after(
            extra,
            ["Métrica", "Valor"],
            [
                ["Documentos / páginas", f"1 / {e['n_paginas']}"],
                ["Hallazgos / cajas tachadas", f"{e['n_hallazgos']} / {e['n_cajas']}"],
                ["Tiempo de procesamiento", f"{e['segundos']:.1f} s"],
                ["Página a revisión (manuscrito)", f"10 (OCR {_fmt(e['ocr'].get('minimo', 40.2), 1)})"],
                ["Recall capa texto NOMBRE (tachados/presentes)",
                 f"{_fmt(e['rec'].get('NOMBRE', 0.8984))} ({nom.get('tachados', 115)}/{nom.get('presentes', 128)})"],
                ["Recall capa texto CÉDULA",
                 f"{_fmt(e['rec'].get('CEDULA', 1.0))} ({ced.get('tachados', 10)}/{ced.get('presentes', 10)})"],
                ["Recall capa texto HC",
                 f"{_fmt(e['rec'].get('HC', 1.0))} ({hc.get('tachados', 12)}/{hc.get('presentes', 12)})"],
                ["Criterio de bloqueo NOMBRE+CÉDULA+HC", "No cumplido (falla NOMBRE)"],
                ["Oraciones extraídas / omitidas por residuo",
                 f"{e['oraciones'] or '—'} / {e['residuos']}"],
            ],
            "Tabla 8 Anonimización OCR y recall de capa de texto (historia de 20 páginas)",
        )
        after_tbl = _insert_blocks_after_element(
            extra._parent,
            tbl8._tbl,
            [
                (
                    "Normal",
                    "Las figuras 13 a 15 resumen la evidencia gráfica de S11: la comparación "
                    "mock frente a LLM real, la distribución de hallazgos OCR por etiqueta y el "
                    "recall de capa de texto. Ninguna figura contiene texto clínico.",
                )
            ],
            rpr,
        )
        if after_tbl:
            last_fig = _insert_picture_after(
                after_tbl,
                CAP / "tripartita_mock_vs_real.png",
                "Figura 13 Comparación tripartita mock frente a LLM real (MEDEC, n=400).",
            )
            last_fig = _insert_picture_after(
                last_fig,
                CAP / "hallazgos_por_etiqueta.png",
                "Figura 14 Hallazgos de de-identificación por etiqueta (agregados, sin texto).",
            )
            _insert_picture_after(
                last_fig,
                CAP / "recall_deidentificacion.png",
                "Figura 15 Recall de de-identificación en la capa de texto del PDF tachado.",
            )


def _ampliar_etica(doc, e, rpr) -> None:
    if _find_paragraph(doc, "Protocolo de de-identificación por capas"):
        return
    _insert_paragraphs_after(
        doc,
        "mejora de la calidad asistencial y la seguridad del paciente.",
        [
            ("List Paragraph", "Protocolo de de-identificación por capas"),
            (
                "Normal",
                "El proyecto trata dos fuentes de naturaleza jurídica distinta. MEDEC es un "
                "conjunto público, en inglés, sin datos identificables, y sobre él se publican "
                "métricas y se admite una API externa. El material de CITIMED son historias "
                "clínicas reales en español: la Ley Orgánica de Protección de Datos Personales "
                "las clasifica como datos sensibles. La regla sin excepciones es que ninguna "
                "nota de CITIMED, original o de-identificada, se transmite a un servicio de "
                "inferencia externo. Cuando se requiera un modelo de lenguaje sobre ese "
                "material, la vía es Ollama local.",
            ),
            (
                "Normal",
                "La de-identificación se ejecuta por completo en el equipo del investigador. "
                "Como referencia metodológica —no jurisdiccional— se adoptan los dieciocho "
                "identificadores del método Safe Harbor de HIPAA, adaptados a Ecuador: cédula "
                "con dígito verificador módulo 10, RUC, telefonía nacional, HCU y edad ≥ 90 "
                "años. Se redactan también los nombres del personal sanitario. Se redacta la "
                "fecha de nacimiento y se conservan, por defecto, las fechas de atención, "
                "desviación consciente del Safe Harbor estricto que se declara y se compensa "
                "con custodia local. El repositorio ignora historias/, salidas/ y los CSV con "
                "texto. Solo salen agregados sin texto.",
            ),
            (
                "Normal",
                f"El tamizado de la capa de texto halló {e['residuos']} residuos de categoría "
                "NOMBRE; esas oraciones se omitieron del corpus y el criterio de bloqueo no se "
                "tiene por cumplido. El límite más relevante se declara sin eufemismos: "
                "Tesseract lee mal el manuscrito, y un nombre escrito a mano fuera de un campo "
                "etiquetado puede escaparse. Por eso la página 10 está marcada a revisión "
                "reforzada y ninguna cifra de este informe debe leerse como garantía de "
                "de-identificación completa. La autorización institucional de CITIMED para uso "
                "académico está declarada por el equipo; el acta formal está pendiente de "
                "adjuntar. Detalle: s11/docs/anexo_etico.md y s11/docs/guia_anotacion.md.",
            ),
        ],
        rpr,
    )


def _conclusiones_y_trazabilidad(doc, e, rpr) -> None:
    p = _find_paragraph(doc, "Viabilidad arquitectónica y protección de datos sanitarios")
    if p and "anonimizador OCR" not in p.text:
        _replace_paragraph_text(
            p,
            "Viabilidad arquitectónica y protección de datos sanitarios: el prototipo opera "
            "en local (FastAPI + React + SQLite) y el pipeline de de-identificación OCR "
            "procesa PDF escaneados sin egreso a internet. El .gitignore excluye historias "
            "originales, PNG de revisión e informes con texto en claro. La evaluación con "
            "LLM real se restringió a MEDEC; para CITIMED se mantiene Ollama on-premise. "
            "Queda pendiente completar la revisión visual dual del lote de 20 páginas, "
            "cerrar el ciclo de corrección de los residuos de NOMBRE y adjuntar el acta "
            "institucional al expediente.",
        )

    if _find_paragraph(doc, "Trazabilidad frente a la entrega S10"):
        return
    ancla = "fomentará la confianza en el sistema bajo el marco de supervisión human-in-the-loop."
    if _find_paragraph(doc, ancla) is None:
        ancla = "Capacitación y adopción tecnológica"
    last = _insert_paragraphs_after(
        doc,
        ancla,
        [
            ("List Paragraph", "Trazabilidad frente a la entrega S10"),
            (
                "Normal",
                "La tabla replica el formato que el evaluador destacó en S10. Lo que sigue "
                "en curso se declara como tal: el kappa 1,0 del piloto es aplicación "
                "determinista de la guía, no doble ciego humano independiente."
            ),
        ],
        rpr,
    )
    if last:
        kappa_estado = (
            f"{e['oraciones'] or '—'} oraciones; kappa piloto {_fmt(e['kappa'])} "
            "(determinista, no doble ciego humano)"
            if e["kappa"] is not None
            else f"{e['oraciones'] or '—'} oraciones extraídas; kappa humano pendiente"
        )
        _insert_table_after(
            last,
            ["Observación S10", "Acción S11", "Evidencia", "Estado"],
            [
                [
                    "pytest no estaba declarado en requirements",
                    "requirements-dev.txt + pyproject.toml; suite unitaria unificada",
                    "21 casos Python en verde; 14/14 frontend; pytest 9.1.1",
                    "Cerrada",
                ],
                [
                    "Evaluación LLM solo en mock",
                    f"eval_llm_real.py con {e['modelo']} sobre {e['n_oraciones']} oraciones MEDEC",
                    f"ROC-AUC TF-IDF {_fmt(e['tfidf'].get('roc_auc', 0.9475))} vs LLM "
                    f"{_fmt(e['zero'].get('roc_auc', 0.5094))}; costo y latencia medidos",
                    "Cerrada",
                ],
                [
                    "Corpus CITIMED y aval ético pendientes",
                    "Anonimizador OCR 1.5 + PDF buscable + anexo ético + guía de anotación",
                    f"{e['n_paginas']} páginas, {e['n_hallazgos']} hallazgos; {kappa_estado}",
                    "Cerrada como muestra piloto (n=100); revisión visual dual pendiente",
                ],
                [
                    "Prototipo sin persistencia ni UI institucional",
                    "Historial SQLite y rediseño React de tres zonas",
                    "api/db.py; s11/capture_prototipo.py; 14 tests de interfaz",
                    "Cerrada",
                ],
                [
                    "Riesgo de filtrar PHI al repositorio",
                    ".gitignore de historias/, salidas/ y ZIP; sanitizador de agregados",
                    "s11/sanitizar_evidencias.py; agregados y hashes sin texto",
                    "Cerrada",
                ],
            ],
            "Tabla 9 Trazabilidad S10 → S11 (observación, acción, evidencia y estado)",
        )

    p_fut = _find_paragraph(doc, "Aunque las pruebas piloto iniciales y la transferencia sobre conjuntos controlados")
    if p_fut:
        _replace_paragraph_text(
            p_fut,
            "Adaptación al contexto clínico local: el OCR sobre una historia real de 20 páginas "
            f"produjo el material de-identificado del que se extrajeron {e['oraciones'] or '—'} "
            "oraciones (muestra piloto, por debajo de las 500–1.000 sugeridas en S7). Ya se "
            f"midió la transferencia MEDEC→CITIMED (ROC-AUC {_fmt(e['cross_auc']) if e.get('cross_auc') is not None else '—'}, "
            f"AUPRC {_fmt(e['cross_auprc']) if e.get('cross_auprc') is not None else '—'}, n=100, "
            "intervalo amplio). El paso inmediato es la revisión visual dual, el doble ciego "
            "humano sobre el 30 % de solape y, con un lote mayor, el fine-tuning en español "
            "odontológico.",
        )


def _anexos(doc, rpr) -> None:
    if _find_paragraph(doc, "Anexo 5. Prototipo React"):
        return
    last = _insert_paragraphs_after(
        doc,
        "Anexo 4. Front End levantado con fastApi.",
        [
            ("Heading 2", "Anexo 5. Prototipo React — vista inicial (S11)"),
            (
                "Normal",
                "Captura del frontend rediseñado (tema claro por defecto). Reproducción: "
                "uvicorn api.main:app --port 8000 y cd frontend && npm run dev; luego "
                "python s11/capture_prototipo.py.",
            ),
        ],
        rpr,
    )
    if last:
        last = _insert_picture_after(
            last, CAP / "prototipo_inicio.png",
            "Figura 9 Prototipo React — pantalla inicial (S11).",
        )
    last = _insert_paragraphs_after(
        doc,
        "Figura 9 Prototipo React",
        [
            ("Heading 2", "Anexo 6. Prototipo React — hallazgo localizado"),
            (
                "Normal",
                "El sistema señala la oración concreta (ejemplo de medicación contraindicada) "
                "y muestra los scores por brazo. No emite un veredicto global sobre la nota.",
            ),
        ],
        rpr,
    )
    if last:
        last = _insert_picture_after(
            last, CAP / "prototipo_hallazgo.png",
            "Figura 10 Prototipo React — hallazgo localizado (S11).",
        )
    last = _insert_paragraphs_after(
        doc,
        "Figura 10 Prototipo React",
        [
            ("Heading 2", "Anexo 7. Historial persistente y tema oscuro"),
            (
                "Normal",
                "El historial SQLite permite volver sobre análisis previos. El tema oscuro es "
                "un atajo no anunciado (cinco pulsaciones sobre el logotipo) y se documenta "
                "porque forma parte del código entregado, no como funcionalidad clínica.",
            ),
        ],
        rpr,
    )
    if last:
        last = _insert_picture_after(
            last, CAP / "prototipo_historial.png",
            "Figura 11 Prototipo React — panel de historial (S11).",
        )
        last = _insert_picture_after(
            last, CAP / "prototipo_tema_oscuro.png",
            "Figura 12 Prototipo React — tema oscuro (S11).",
        )
    ancla_8 = "Figura 12 Prototipo React" if _find_paragraph(doc, "Figura 12 Prototipo React") else "Anexo 7. Historial persistente"
    _insert_paragraphs_after(
        doc,
        ancla_8,
        [
            ("Heading 2", "Anexo 8. Documentos éticos y de anotación"),
            (
                "Normal",
                "El detalle normativo, la correspondencia con los dieciocho identificadores "
                "tipo HIPAA, la cadena de custodia y los marcadores pendientes del acta "
                "institucional están en s11/docs/anexo_etico.md. Los criterios de etiquetado "
                "y el procedimiento de doble anotación están en s11/docs/guia_anotacion.md. "
                "El protocolo de verificación humana está en "
                f"s11/docs/protocolo_verificacion_humana.md. Repositorio: {REPO_URL}.",
            ),
        ],
        rpr,
    )


def export_pdf(docx_path: Path | None = None) -> None:
    src = docx_path or OUT_DOCX
    pdf = DOCS / "Proyecto_Capstone_S11.pdf"
    try:
        import comtypes.client  # type: ignore

        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        d = word.Documents.Open(str(src.resolve()))
        d.SaveAs(str(pdf.resolve()), FileFormat=17)
        d.Close()
        word.Quit()
        print(f"[ok] PDF (MS Word) -> {pdf}")
        return
    except Exception as exc:
        print(f"[info] Word COM no disponible: {exc}")
    print(f"[warn] Exporte a mano {src} → PDF para Moodle/Teams.")


def main() -> None:
    import importlib.util

    png = CAP / "flujo_prototipo.png"
    if not png.exists():
        spec = importlib.util.spec_from_file_location(
            "generar_flujo_prototipo", S11 / "generar_flujo_prototipo.py"
        )
        mod = importlib.util.module_from_spec(spec)
        assert spec.loader is not None
        spec.loader.exec_module(mod)
        mod.main()

    ev = load_evidence()
    if not ev["llm"]:
        print("[warn] Falta s11/evidencias/metricas_llm_real.json")
    saved = actualizar(ev)
    export_pdf(saved)


if __name__ == "__main__":
    sys.exit(main() or 0)
